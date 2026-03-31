"""
Document Text Extraction
========================
Extracts plain text from PDF, Word, and Excel files.
"""

import io


DOC_CHAR_LIMIT = 50_000
TOTAL_DOC_CHAR_LIMIT = 150_000


def extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts)


def extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_xlsx(file_bytes: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(v) for v in row if v is not None)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    name = filename.lower()
    try:
        if name.endswith(".pdf"):
            return extract_pdf(file_bytes)
        elif name.endswith(".docx"):
            return extract_docx(file_bytes)
        elif name.endswith((".xlsx", ".xls")):
            return extract_xlsx(file_bytes)
        else:
            return ""
    except Exception as exc:
        return f"[Extraction error: {exc}]"
