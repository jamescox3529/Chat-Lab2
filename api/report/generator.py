"""
Report Document Generator
=========================
Produces branded .docx or .pdf from AI-generated markdown report content.

Word:  python-docx (already in requirements)
PDF:   reportlab  (add: reportlab>=4.0.0 to requirements.txt)
"""
from __future__ import annotations

import io
import os
import re

# ---------------------------------------------------------------------------
# python-docx
# ---------------------------------------------------------------------------
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# reportlab
# ---------------------------------------------------------------------------
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, black
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
    PageBreak,
    NextPageTemplate,
    Image,
    KeepTogether,
)
from reportlab.platypus.flowables import HRFlowable  # noqa: F811
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader

# ---------------------------------------------------------------------------
# Brand colours
# ---------------------------------------------------------------------------
TEAL_RGB = RGBColor(0x4A, 0x8B, 0x8C)
NEAR_BLACK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
GREY_RGB = RGBColor(0x80, 0x80, 0x80)

TEAL_RL = HexColor("#4A8B8C")
NEAR_BLACK_RL = HexColor("#1A1A1A")
GREY_RL = HexColor("#808080")

# ---------------------------------------------------------------------------
# Asset paths
# ---------------------------------------------------------------------------
_HERE = os.path.dirname(os.path.abspath(__file__))
_ASSETS = os.path.normpath(os.path.join(_HERE, "..", "assets"))

_LOGO_CANDIDATES = [
    os.path.join(_ASSETS, "logo.jpg"),
    # local dev fallbacks
    os.path.normpath(os.path.join(_HERE, "..", "..", "roundtable-logo light mode.jpg")),
    os.path.normpath(os.path.join(_HERE, "..", "..", "web", "public", "icon-light.png")),
]

_FONTS_DIR = os.path.join(_ASSETS, "fonts")
_POPPINS_BOLD    = os.path.join(_FONTS_DIR, "Poppins-Bold.ttf")
_POPPINS_REGULAR = os.path.join(_FONTS_DIR, "Poppins-Regular.ttf")


def _get_logo_path() -> str | None:
    for p in _LOGO_CANDIDATES:
        if os.path.exists(p):
            return p
    return None


def _register_heading_font() -> str:
    """Register Poppins-Bold if available; return the font name to use for headings."""
    if os.path.exists(_POPPINS_BOLD):
        try:
            pdfmetrics.registerFont(TTFont("Poppins-Bold", _POPPINS_BOLD))
            return "Poppins-Bold"
        except Exception:
            pass
    return "Helvetica-Bold"


def _logo_dimensions(logo_path: str, display_width_cm: float) -> tuple[float, float]:
    """Return (width, height) in reportlab points preserving the image's aspect ratio."""
    try:
        ir = ImageReader(logo_path)
        nat_w, nat_h = ir.getSize()
        w = display_width_cm * cm
        h = w * (nat_h / nat_w)
        return w, h
    except Exception:
        return display_width_cm * cm, display_width_cm * cm  # square fallback


# ===========================================================================
# Shared markdown parser
# ===========================================================================

def _parse_md(md: str) -> list[tuple[str, str]]:
    """
    Parse AI markdown into (kind, text) tuples.
    Kinds: h1  h2  bullet  para  empty
    """
    elements: list[tuple[str, str]] = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if line.startswith("## "):
            elements.append(("h1", line[3:].strip()))
        elif line.startswith("### "):
            elements.append(("h2", line[4:].strip()))
        elif line.startswith("- ") or line.startswith("* "):
            elements.append(("bullet", line[2:].strip()))
        elif line.startswith("  - ") or line.startswith("  * "):
            elements.append(("bullet", line[4:].strip()))
        elif line == "":
            elements.append(("empty", ""))
        else:
            elements.append(("para", line.lstrip("#").strip()))
    return elements


def _split_bold(text: str) -> list[tuple[str, bool]]:
    """Split on **..** → [(segment, is_bold), ...]"""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    return [(p, i % 2 == 1) for i, p in enumerate(parts) if p]


# ===========================================================================
# Word (.docx) generation
# ===========================================================================

def _set_run(run, size_pt: float, bold: bool = False,
             color: RGBColor | None = None, italic: bool = False,
             font_name: str = "Arial") -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_para(doc: Document, text: str, size_pt: float = 11,
              bold: bool = False, color: RGBColor | None = None,
              align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before: float = 0, space_after: float = 4,
              font_name: str = "Arial") -> object:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    for seg, is_bold in _split_bold(text):
        run = p.add_run(seg)
        _set_run(run, size_pt, bold=(bold or is_bold), color=color, font_name=font_name)
    return p


def _teal_rule(doc: Document) -> None:
    """Thin teal bottom border on an empty paragraph — used below H1."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")       # 0.75 pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4A8B8C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_num_field(paragraph) -> None:
    """Insert 'Page <PAGE>' into paragraph using a Word field."""
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = NEAR_BLACK_RGB

    def _field_char(type_: str):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), type_)
        r = OxmlElement("w:r")
        r.append(fc)
        return r

    instr_r = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_r.append(instr)

    paragraph._p.append(_field_char("begin"))
    paragraph._p.append(instr_r)
    paragraph._p.append(_field_char("end"))


def generate_docx(
    content: str,
    title: str,
    room_name: str,
    user_name: str,
    date: str,
    subtitle: str = "Consult Room Report",
    participants_label: str = "Prepared for",
) -> bytes:
    # Use Poppins for headings/titles if font is present; fall back to Arial.
    heading_font = "Poppins" if os.path.exists(_POPPINS_BOLD) else "Arial"

    doc = Document()

    # ---- Page setup --------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.different_first_page_header_footer = True

    logo_path = _get_logo_path()

    # ---- Body header: logo right-aligned -----------------------------------
    header = section.header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_para.paragraph_format.space_after = Pt(0)
    if logo_path:
        h_para.add_run().add_picture(logo_path, width=Cm(2.5))
    else:
        run = h_para.add_run("Roundtable")
        _set_run(run, 10, bold=True, color=TEAL_RGB)

    # ---- Body footer: page number right-aligned ----------------------------
    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_num_field(f_para)

    # ---- Cover footer: Private & Confidential ------------------------------
    cf_para = section.first_page_footer.paragraphs[0]
    run = cf_para.add_run("Private & Confidential")
    _set_run(run, 9, color=GREY_RGB)

    # ========================================================================
    # COVER PAGE
    # ========================================================================

    # Logo — top right
    if logo_path:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lp.add_run().add_picture(logo_path, width=Cm(5.0))

    # Whitespace
    for _ in range(7):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)

    # Report title (28pt bold near-black)
    _add_para(doc, title, size_pt=28, bold=True, color=NEAR_BLACK_RGB,
              space_before=0, space_after=6, font_name=heading_font)

    # Subtitle (13pt teal)
    _add_para(doc, subtitle, size_pt=13, color=TEAL_RGB,
              space_before=0, space_after=4, font_name=heading_font)

    # Teal rule
    _teal_rule(doc)

    # Metadata block
    for line in [
        f"Room: {room_name}",
        f"{participants_label}: {user_name}",
        f"Date: {date}",
    ]:
        _add_para(doc, line, size_pt=11, color=NEAR_BLACK_RGB,
                  space_before=2, space_after=2)

    # Page break → body
    doc.add_page_break()

    # ========================================================================
    # REPORT BODY
    # ========================================================================

    for kind, text in _parse_md(content):
        if kind == "h1":
            _add_para(doc, text, size_pt=13, bold=True, color=TEAL_RGB,
                      space_before=14, space_after=0, font_name=heading_font)
            _teal_rule(doc)
        elif kind == "h2":
            _add_para(doc, text, size_pt=11, bold=True, color=TEAL_RGB,
                      space_before=8, space_after=2, font_name=heading_font)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            for seg, is_bold in _split_bold(text):
                run = p.add_run(seg)
                _set_run(run, 11, bold=is_bold, color=NEAR_BLACK_RGB)
        elif kind == "para" and text:
            _add_para(doc, text, size_pt=11, color=NEAR_BLACK_RGB,
                      space_before=0, space_after=4)
        # empty → skip

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# PDF generation (reportlab)
# ===========================================================================

def _rl_styles(heading_font: str) -> dict[str, ParagraphStyle]:
    """Build the reportlab paragraph styles.
    Body copy, bullets, and metadata use Helvetica; only headings use heading_font.
    """
    return {
        "body": ParagraphStyle("body", fontName="Helvetica", fontSize=11,
                               leading=11 * 1.15, textColor=NEAR_BLACK_RL,
                               spaceAfter=6),
        "h1": ParagraphStyle("h1", fontName=heading_font, fontSize=13,
                             leading=13 * 1.2, textColor=TEAL_RL,
                             spaceBefore=14, spaceAfter=2),
        "h2": ParagraphStyle("h2", fontName=heading_font, fontSize=11,
                             leading=11 * 1.2, textColor=TEAL_RL,
                             spaceBefore=8, spaceAfter=2),
        "bullet": ParagraphStyle("bullet", fontName="Helvetica", fontSize=11,
                                 leading=11 * 1.15, textColor=NEAR_BLACK_RL,
                                 leftIndent=18, firstLineIndent=0,
                                 bulletIndent=6, spaceAfter=3),
        "cover_title": ParagraphStyle("cover_title", fontName=heading_font,
                                      fontSize=28, leading=32,
                                      textColor=NEAR_BLACK_RL, spaceAfter=8),
        "cover_sub": ParagraphStyle("cover_sub", fontName=heading_font,
                                    fontSize=13, leading=16,
                                    textColor=TEAL_RL, spaceAfter=6),
        "cover_meta": ParagraphStyle("cover_meta", fontName="Helvetica",
                                     fontSize=11, leading=14,
                                     textColor=NEAR_BLACK_RL, spaceAfter=3),
    }


def _md_to_rl_html(text: str) -> str:
    """Convert **bold** markers to reportlab HTML bold tags."""
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)


class _PDFDoc(BaseDocTemplate):
    """Thin subclass to carry logo_path into page callbacks."""
    logo_path: str | None = None


def build_debate_markdown(
    rounds: list[dict],
    synthesis: str,
) -> str:
    """
    Assemble a debate transcript + synthesis into a single markdown string
    compatible with _parse_md().  Each round becomes an h1 section; each
    persona within a round becomes an h2.  The synthesis follows with its
    own h1 sub-sections (Recommendation, Key Reasoning, etc.).
    """
    parts: list[str] = []

    for round_data in rounds:
        parts.append(f"## Round {round_data['round']}: {round_data['label']}")
        for response in round_data.get("responses", []):
            parts.append(f"### {response['role']}")
            parts.append(response["content"])
            parts.append("")

    parts.append("## Panel Verdict")
    parts.append("")
    parts.append(synthesis)

    return "\n\n".join(parts)


def generate_pdf(
    content: str,
    title: str,
    room_name: str,
    user_name: str,
    date: str,
    subtitle: str = "Consult Room Report",
    participants_label: str = "Prepared for",
) -> bytes:
    buf = io.BytesIO()
    logo_path = _get_logo_path()
    heading_font = _register_heading_font()

    pw, ph = A4
    margin = 2.5 * cm
    text_w = pw - 2 * margin
    text_h = ph - 2 * margin

    doc = _PDFDoc(buf, pagesize=A4,
                  leftMargin=margin, rightMargin=margin,
                  topMargin=margin, bottomMargin=margin)
    doc.logo_path = logo_path

    # ---- Page templates ----------------------------------------------------

    cover_frame = Frame(margin, margin, text_w, text_h, id="cover",
                        leftPadding=0, rightPadding=0,
                        topPadding=0, bottomPadding=0)

    body_top_margin = margin + 1.5 * cm   # room for header
    body_frame = Frame(margin, margin, text_w,
                       ph - body_top_margin - margin,
                       id="body",
                       leftPadding=0, rightPadding=0,
                       topPadding=0, bottomPadding=0)

    def _on_cover(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(GREY_RL)
        canvas.drawString(margin, margin * 0.7, "Private & Confidential")
        canvas.restoreState()

    def _on_body(canvas, doc_):
        canvas.saveState()
        # Header: logo right-aligned, aspect-ratio correct
        if doc_.logo_path and os.path.exists(doc_.logo_path):
            logo_w, logo_h = _logo_dimensions(doc_.logo_path, 2.5)
            canvas.drawImage(
                doc_.logo_path,
                pw - margin - logo_w,
                ph - margin * 0.85 - logo_h,
                width=logo_w, height=logo_h,
            )
        else:
            canvas.setFont(heading_font, 10)
            canvas.setFillColor(TEAL_RL)
            canvas.drawRightString(pw - margin, ph - margin * 0.75, "Roundtable")

        # Footer: Page N
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(NEAR_BLACK_RL)
        canvas.drawRightString(pw - margin, margin * 0.7, f"Page {doc_.page}")
        canvas.restoreState()

    doc.addPageTemplates([
        PageTemplate(id="Cover", frames=[cover_frame], onPage=_on_cover),
        PageTemplate(id="Body",  frames=[body_frame],  onPage=_on_body),
    ])

    styles = _rl_styles(heading_font)
    story = []

    # ========================================================================
    # COVER PAGE
    # ========================================================================

    story.append(NextPageTemplate("Cover"))

    # Logo — right-aligned, aspect-ratio correct
    if logo_path and os.path.exists(logo_path):
        logo_w, logo_h = _logo_dimensions(logo_path, 5.0)
        img = Image(logo_path, width=logo_w, height=logo_h)
        img.hAlign = "RIGHT"
        story.append(img)
    else:
        story.append(Paragraph("<b>Roundtable</b>",
                                ParagraphStyle("logo_text", fontName=heading_font,
                                               fontSize=16, textColor=TEAL_RL,
                                               alignment=TA_RIGHT)))

    # Whitespace before title
    story.append(Spacer(1, 5 * cm))

    # Title
    story.append(Paragraph(title, styles["cover_title"]))

    # Subtitle
    story.append(Paragraph(subtitle, styles["cover_sub"]))

    # Teal rule
    story.append(HRFlowable(width="100%", thickness=1, color=TEAL_RL,
                             spaceAfter=8))

    # Metadata
    for line in [
        f"Room: {room_name}",
        f"{participants_label}: {user_name}",
        f"Date: {date}",
    ]:
        story.append(Paragraph(line, styles["cover_meta"]))

    # Switch template + page break
    story.append(NextPageTemplate("Body"))
    story.append(PageBreak())

    # ========================================================================
    # REPORT BODY
    # ========================================================================

    for kind, text in _parse_md(content):
        html = _md_to_rl_html(text)
        if kind == "h1":
            story.append(KeepTogether([
                Paragraph(html, styles["h1"]),
                HRFlowable(width="100%", thickness=1, color=TEAL_RL,
                           spaceBefore=2, spaceAfter=6),
            ]))
        elif kind == "h2":
            story.append(Paragraph(html, styles["h2"]))
        elif kind == "bullet":
            story.append(Paragraph(f"\u2022&nbsp;&nbsp;{html}", styles["bullet"]))
        elif kind == "para" and text:
            story.append(Paragraph(html, styles["body"]))
        # empty → skip

    doc.build(story)
    return buf.getvalue()
